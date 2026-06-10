# -*- coding: utf-8 -*-
"""
严格按目录结构重新生成 LaTeX → PDF → Word 文档。
"""
import re
import sys
import io
import subprocess
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

INPUT_FILE = Path("outputs/定本_完整处理/定本 柄谷行人文学论集 (柄谷行人 著，陈言 译) (z-library.sk, 1lib.sk, z-lib.sk)_20260609_155715/translation.md")
OUT_DIR = Path("outputs/定本_完整处理/定本 柄谷行人文学论集 (柄谷行人 著，陈言 译) (z-library.sk, 1lib.sk, z-lib.sk)_20260609_155715")

# ── 目录定义（严格按原书目录） ──────────────────────────────────
# 每个条目: (标题, 父级, 关键词用于匹配翻译文本)
# 父级: None=顶级, "part1"=第一部, "part2"=第二部
TOC = [
    ("序 文", None, ["序文"]),
    ("Ⅰ", None, []),
    ("《亚历山大四重奏》的辩证法", "part1", ["亚历山大四重奏"]),
    ("漱石试论——意识与自然", "part1", ["漱石试论"]),
    ("意义的病——麦克白论", "part1", ["麦克白"]),
    ("历史与自然——森鸥外论", "part1", ["森鸥外"]),
    ("关于坂口安吾的《日本文化私观》", "part1", ["坂口安吾"]),
    ("关于历史——武田泰淳", "part1", ["武田泰淳"]),
    ("Ⅱ", None, []),
    ("漱石的多样性", "part2", ["漱石的多样性"]),
    ("坂口安吾，其可能性的中心", "part2", ["坂口安吾，其可能性"]),
    ("梦的世界——岛尾敏雄", "part2", ["梦的世界", "岛尾敏雄"]),
    ("中上健次与福克纳", "part2", ["中上健次"]),
    ("翻译家四迷", "part2", ["翻译家四迷", "二叶亭四迷"]),
    ("文学的衰灭", "part2", ["文学的衰灭"]),
    ("初刊·底本一览", None, ["初刊", "底本"]),
    ("译后记", None, ["译后记"]),
    ("柄谷行人：移动的文学批评", None, ["移动的文学批评"]),
]


def load_translation(path: Path) -> str:
    content = path.read_text(encoding="utf-8")
    # 去掉页面标记
    content = re.sub(r'^## 第 \d+ 页\s*$', '', content, flags=re.MULTILINE)
    # 去掉页眉页脚（如 "004 | 定本柄谷行人文学论集"）
    content = re.sub(r'^\d{3,4}\s*\|.*$', '', content, flags=re.MULTILINE)
    # 去掉页码标注格式（如 "翻译家四迷 | 327"）
    content = re.sub(r'^.+\|\s*\d+$', '', content, flags=re.MULTILINE)
    # 去掉孤立页码
    content = re.sub(r'^第 \d+ 页$', '', content, flags=re.MULTILINE)
    # 去掉单独的数字页码
    content = re.sub(r'^\d{1,4}$', '', content, flags=re.MULTILINE)
    return content


def find_section_positions(text: str) -> list[tuple[str, int]]:
    """在文本中找到各章节的起始位置（跳过目录区域）"""
    positions = []
    lines = text.split("\n")

    # 先找到目录区域的范围（跳过）
    toc_start = None
    toc_end = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if re.match(r'^《亚历山大四重奏》的辩证法\s*/\s*\d+$', stripped):
            toc_start = i
        if toc_start and re.match(r'^柄谷行人：移动的文学批评\s*/\s*\d+$', stripped):
            toc_end = i + 2  # 跳过目录区域
            break

    for i, line in enumerate(lines):
        # 跳过目录区域
        if toc_start and toc_end and toc_start <= i <= toc_end:
            continue

        stripped = line.strip()
        for title, parent, keywords in TOC:
            if not keywords:
                continue
            for kw in keywords:
                # 精确匹配：标题行应该完全等于标题，或者是独立的标题行
                # 排除在正文中提到标题的情况
                if stripped == title:
                    # 计算字符偏移
                    offset = sum(len(l) + 1 for l in lines[:i])
                    positions.append((title, offset))
                    break

    # 去重并排序，保留每个标题第一次出现的位置
    seen = set()
    unique = []
    for title, offset in positions:
        if title not in seen:
            seen.add(title)
            unique.append((title, offset))
    unique.sort(key=lambda x: x[1])
    return unique


def split_by_sections(text: str) -> list[tuple[str, str]]:
    """按章节标题分割文本"""
    positions = find_section_positions(text)
    sections = []

    # 找到"序文"之前的前言部分（封面、版权页等）
    if positions:
        first_offset = positions[0][1]
        preamble = text[:first_offset].strip()
        if preamble:
            sections.append(("前言", preamble))

    for i, (title, offset) in enumerate(positions):
        if i + 1 < len(positions):
            next_offset = positions[i + 1][1]
            content = text[offset:next_offset].strip()
        else:
            content = text[offset:].strip()
        sections.append((title, content))

    return sections


def clean_text(text: str) -> str:
    """清理文本：去除噪声行"""
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        # 跳过噪声行
        if re.match(r'^ISBN \d', stripped):
            continue
        if re.match(r'^定价[：:]', stripped):
            continue
        if re.match(r'^微信扫描', stripped):
            continue
        if re.match(r'^上架建议', stripped):
            continue
        if re.match(r'^9 \d{10,}', stripped):
            continue
        if re.match(r'^新浪微博', stripped):
            continue
        if re.match(r'^微\s*信[：:]', stripped):
            continue
        if re.match(r'^淘宝店铺', stripped):
            continue
        if re.match(r'^本社常年法律顾问', stripped):
            continue
        if re.match(r'^凡有印装质量问题', stripped):
            continue
        if re.match(r'^责任编辑', stripped):
            continue
        if re.match(r'^责任印制', stripped):
            continue
        if re.match(r'^出版发行', stripped):
            continue
        if re.match(r'^地\s+址[：:]', stripped):
            continue
        if re.match(r'^电\s+话[：:]', stripped):
            continue
        if re.match(r'^传\s+真[：:]', stripped):
            continue
        if re.match(r'^经\s+销[：:]', stripped):
            continue
        if re.match(r'^印\s+刷[：:]', stripped):
            continue
        if re.match(r'^开\s+本[：:]', stripped):
            continue
        if re.match(r'^字\s+数[：:]', stripped):
            continue
        if re.match(r'^印\s+张[：:]', stripped):
            continue
        if re.match(r'^版\s+次[：:]', stripped):
            continue
        if re.match(r'^印\s+次[：:]', stripped):
            continue
        if re.match(r'^ICCTP', stripped):
            continue
        if re.match(r'^Central Compilation', stripped):
            continue
        if re.match(r'^著$', stripped):
            continue
        if re.match(r'^译$', stripped):
            continue
        if re.match(r'^陈言\s*译$', stripped):
            continue
        if re.match(r'^柄谷行人\s*著$', stripped):
            continue
        if re.match(r'^中央编译出版社', stripped):
            continue
        if re.match(r'^定本$', stripped):
            continue
        if re.match(r'^柄谷行人文学论集$', stripped):
            continue
        if re.match(r'^柄谷行人$', stripped):
            continue
        if re.match(r'^图书在版编目', stripped):
            continue
        if re.match(r'^著作权登记号', stripped):
            continue
        if re.match(r'^TEIHON', stripped):
            continue
        if re.match(r'^by Kojin', stripped):
            continue
        if re.match(r'^Originally published', stripped):
            continue
        if re.match(r'^This simplified', stripped):
            continue
        if re.match(r'^by arrangement', stripped):
            continue
        if re.match(r'^\d+页$', stripped):
            continue
        if re.match(r'^定本柄谷行人文文学论集 \[存疑\]', stripped):
            continue
        if re.match(r'^作者简介$', stripped):
            continue
        if re.match(r'^译后记$', stripped):
            continue
        if re.match(r'^柄谷行人：移动的文学批评$', stripped):
            continue
        # 跳过页眉页脚格式：标题 | 页码
        if re.match(r'^.+\|\s*\d+$', stripped):
            continue
        # 跳过单独的页码行
        if re.match(r'^\d{1,4}$', stripped):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def escape_latex(text: str) -> str:
    """转义 LaTeX 特殊字符"""
    # 保留脚注标记
    text = text.replace('①', '\\textsuperscript{①}')
    text = text.replace('②', '\\textsuperscript{②}')
    text = text.replace('③', '\\textsuperscript{③}')
    text = text.replace('④', '\\textsuperscript{④}')
    text = text.replace('⑤', '\\textsuperscript{⑤}')
    text = text.replace('⑥', '\\textsuperscript{⑥}')
    text = text.replace('⑦', '\\textsuperscript{⑦}')
    text = text.replace('⑧', '\\textsuperscript{⑧}')
    text = text.replace('⑨', '\\textsuperscript{⑨}')
    text = text.replace('⑩', '\\textsuperscript{⑩}')
    # 转义特殊字符
    text = text.replace('&', '\\&')
    text = text.replace('%', '\\%')
    text = text.replace('$', '\\$')
    text = text.replace('#', '\\#')
    text = text.replace('_', '\\_')
    text = text.replace('{', '\\{')
    text = text.replace('}', '\\}')
    text = text.replace('~', '\\textasciitilde{}')
    text = text.replace('^', '\\textasciicircum{}')
    # 保留引号
    text = text.replace('"', '``')
    text = text.replace('"', "''")
    text = text.replace('「', "``")
    text = text.replace('」', "''")
    text = text.replace('『', "``")
    text = text.replace('』', "''")
    return text


def text_to_paragraphs(text: str) -> list[str]:
    """将文本按空行分割为段落"""
    paragraphs = []
    current = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            if current:
                paragraphs.append(" ".join(current))
                current = []
        else:
            current.append(stripped)
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs


def build_latex(sections: list[tuple[str, str]]) -> str:
    """构建 LaTeX 文档"""
    # LaTeX 文档头
    preamble = r"""
\documentclass[a4paper,12pt]{book}
\usepackage[UTF8]{ctex}
\usepackage{fontspec}
\usepackage{geometry}
\usepackage{fancyhdr}
\usepackage{titlesec}
\usepackage{setspace}
\usepackage{parskip}

% 页面设置
\geometry{
  left=3cm,
  right=2.5cm,
  top=2.5cm,
  bottom=2.5cm
}

% 行距
\onehalfspacing

% 段落间距
\setlength{\parindent}{2em}
\setlength{\parskip}{0.3em}

% 页眉页脚
\pagestyle{fancy}
\fancyhf{}
\fancyhead[LE]{\leftmark}
\fancyhead[RO]{\rightmark}
\fancyfoot[C]{\thepage}

% 标题样式
\titleformat{\chapter}[display]
  {\normalfont\huge\bfseries\centering}
  {\chaptertitlename\ \thechapter}
  {20pt}
  {\Huge}
\titlespacing*{\chapter}{0pt}{50pt}{40pt}

\titleformat{\section}
  {\normalfont\Large\bfseries}
  {\thesection}
  {1em}
  {}
\titlespacing*{\section}{0pt}{3.5ex plus 1ex minus .2ex}{2.3ex plus .2ex}

\titleformat{\subsection}
  {\normalfont\large\bfseries}
  {\thesubsection}
  {1em}
  {}
\titlespacing*{\subsection}{0pt}{3ex plus 1ex minus .2ex}{1.5ex plus .2ex}

\begin{document}

% 标题页
\begin{titlepage}
\centering
\vspace*{3cm}
{\Huge\bfseries 定本 柄谷行人文学论集\par}
\vspace{2cm}
{\Large 柄谷行人 著\par}
\vspace{1cm}
{\Large 陈言 译\par}
\vspace{3cm}
{\large 中央编译出版社\par}
\end{titlepage}

\tableofcontents
\clearpage

"""

    body = []

    for title, content in sections:
        content = clean_text(content)
        paragraphs = text_to_paragraphs(content)

        if title == "前言":
            # 前言部分（封面信息等）跳过或简化
            continue
        elif title == "序 文":
            body.append("\\chapter*{序 文}")
            body.append("\\addcontentsline{toc}{chapter}{序 文}")
            body.append("")
        elif title == "Ⅰ":
            body.append("\\part{Ⅰ}")
            body.append("")
        elif title == "Ⅱ":
            body.append("\\part{Ⅱ}")
            body.append("")
        elif title in ["初刊·底本一览", "译后记", "柄谷行人：移动的文学批评"]:
            body.append(f"\\chapter*{{{escape_latex(title)}}}")
            body.append(f"\\addcontentsline{{toc}}{{chapter}}{{{escape_latex(title)}}}")
            body.append("")
        else:
            # 具体文章标题
            body.append(f"\\section{{{escape_latex(title)}}}")
            body.append("")

        # 添加段落
        for para in paragraphs:
            escaped = escape_latex(para)
            # 检测是否为脚注行
            if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', para):
                body.append(f"\\noindent {escaped}")
            else:
                body.append(escaped)
            body.append("")

    # 文档尾
    footer = r"""
\end{document}
"""

    return preamble + "\n".join(body) + footer


def main():
    print("📖 读取翻译文件...")
    text = load_translation(INPUT_FILE)
    print(f"   总字符数: {len(text):,}")

    print("📐 按目录分割章节...")
    sections = split_by_sections(text)
    print(f"   章节数: {len(sections)}")
    for title, content in sections:
        print(f"     - {title}: {len(content):,} 字符")

    print("📝 生成 LaTeX 文件...")
    latex_content = build_latex(sections)
    latex_file = OUT_DIR / "translation.tex"
    latex_file.write_text(latex_content, encoding="utf-8")
    print(f"   已生成: {latex_file.name}")

    print("🔨 编译 LaTeX → PDF...")
    try:
        # 使用 xelatex 编译（支持中文）
        for pass_num in [1, 2]:
            result = subprocess.run(
                ["xelatex", "-interaction=nonstopmode", "-halt-on-error",
                 "-output-directory", str(OUT_DIR), str(latex_file)],
                capture_output=True, text=True, timeout=300,
                encoding="utf-8", errors="replace",
            )
            if result.returncode != 0 and pass_num == 2:
                print(f"   ⚠ LaTeX 编译警告（第 {pass_num} 次）")
                # 检查是否有严重错误
                if "fatal error" in result.stdout.lower():
                    print("   ❌ LaTeX 编译失败")
                    print(result.stdout[-500:])
                    return
        pdf_file = OUT_DIR / "translation.pdf"
        if pdf_file.exists():
            size = pdf_file.stat().st_size / 1024 / 1024
            print(f"   ✅ PDF 已生成: {pdf_file.name} ({size:.1f} MB)")
        else:
            print("   ❌ PDF 文件未生成")
    except FileNotFoundError:
        print("   ⚠ 未找到 xelatex，跳过 PDF 生成")
        print("   请安装 TeX Live 或 MiKTeX")
    except subprocess.TimeoutExpired:
        print("   ❌ LaTeX 编译超时")

    print("📝 生成 Word 文档...")
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document()

        # 设置默认字体
        styles = doc.styles
        for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = "Microsoft YaHei"
                if style._element.rPr is not None:
                    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

        # 标题页
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("\n\n\n\n\n")
        run.font.size = Pt(12)
        run = title_para.add_run("定本 柄谷行人文学论集")
        run.font.size = Pt(28)
        run.bold = True

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run("\n\n柄谷行人 著")
        run.font.size = Pt(16)

        translator_para = doc.add_paragraph()
        translator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = translator_para.add_run("\n陈言 译")
        run.font.size = Pt(16)

        publisher_para = doc.add_paragraph()
        publisher_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = publisher_para.add_run("\n\n\n\n中央编译出版社")
        run.font.size = Pt(14)

        doc.add_page_break()

        # 目录页
        toc_heading = doc.add_heading("目 录", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        # 添加目录条目
        for title, parent, _ in TOC:
            if parent is None and title not in ["Ⅰ", "Ⅱ"]:
                # 顶级条目
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.font.size = Pt(12)
                run.bold = True
            elif title in ["Ⅰ", "Ⅱ"]:
                # 部标题
                para = doc.add_paragraph()
                run = para.add_run(f"\n{title}")
                run.font.size = Pt(14)
                run.bold = True
            else:
                # 子条目
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
                run = para.add_run(title)
                run.font.size = Pt(11)

        doc.add_page_break()

        # 正文内容
        for title, content in sections:
            content = clean_text(content)
            paragraphs = text_to_paragraphs(content)

            if title == "前言":
                continue
            elif title == "序 文":
                heading = doc.add_heading("序 文", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif title == "Ⅰ":
                doc.add_page_break()
                heading = doc.add_heading("Ⅰ", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif title == "Ⅱ":
                doc.add_page_break()
                heading = doc.add_heading("Ⅱ", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif title in ["初刊·底本一览", "译后记", "柄谷行人：移动的文学批评"]:
                doc.add_page_break()
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title, level=2)

            # 添加段落
            for para_text in paragraphs:
                para = doc.add_paragraph()
                # 处理脚注
                if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', para_text):
                    run = para.add_run(para_text)
                    run.font.size = Pt(9)
                else:
                    run = para.add_run(para_text)
                    run.font.size = Pt(11)
                    para.paragraph_format.first_line_indent = Cm(0.74)

        word_file = OUT_DIR / "translation_final.docx"
        doc.save(word_file)
        size = word_file.stat().st_size / 1024
        print(f"   ✅ Word 已生成: {word_file.name} ({size:.1f} KB)")

    except ImportError:
        print("   ⚠ 未安装 python-docx，跳过 Word 生成")

    print("\n✅ 完成！")


if __name__ == "__main__":
    main()
