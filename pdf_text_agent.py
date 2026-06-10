from __future__ import annotations

import argparse
import base64
import datetime as dt
import html
import json
import mimetypes
import os
import re
import shutil
import sys
import time
import urllib.error
import urllib.request
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

import fitz
from docx import Document
from docx.oxml.ns import qn
from openai import AuthenticationError, OpenAI
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
PDF_EXTENSIONS = {".pdf"}


@dataclass
class PageResult:
    source: str
    kind: str
    page_number: int
    embedded_text: str = ""
    ocr_text: str = ""
    translation: str = ""
    notes: str = ""
    image_path: str = ""

    @property
    def source_text(self) -> str:
        return (self.ocr_text or self.embedded_text or "").strip()


@dataclass(frozen=True)
class ProviderProfile:
    core: str
    api_key_env: str
    base_url_env: str
    model_env: str
    default_base_url: str
    default_model: str
    token_param: str
    default_vision_detail: str = "xhigh"
    supports_json_mode: bool = True
    supports_vision_detail: bool = True
    supports_image_input: bool = True
    api_key_header: str = "authorization"
    context_window_tokens: int | None = None
    default_ocr_max_tokens: int = 12000
    default_translate_max_tokens: int = 12000
    default_summary_max_tokens: int = 8000
    default_translation_chunk_chars: int = 16000
    default_summary_chunk_chars: int = 24000
    thinking: dict[str, str] | None = None
    reasoning_effort: str | None = None
    native_http: bool = False
    api_style: str = "chat"


PROVIDER_PROFILES: dict[str, ProviderProfile] = {
    "seed": ProviderProfile(
        core="seed",
        api_key_env="SEED_API_KEY",
        base_url_env="SEED_BASE_URL",
        model_env="SEED_MODEL",
        default_base_url="https://ark.cn-beijing.volces.com/api/v3/responses",
        default_model="doubao-seed-2-0-lite-260215",
        token_param="max_tokens",
        default_vision_detail="xhigh",
        supports_json_mode=True,
        supports_vision_detail=True,
        api_key_header="authorization",
        native_http=True,
        api_style="responses",
    ),
    "mimo": ProviderProfile(
        core="mimo",
        api_key_env="MIMO_API_KEY",
        base_url_env="MIMO_BASE_URL",
        model_env="MIMO_MODEL",
        default_base_url="https://token-plan-cn.xiaomimimo.com/v1",
        default_model="mimo-v2.5",
        token_param="max_completion_tokens",
        default_vision_detail="high",
        supports_json_mode=True,
        supports_vision_detail=False,
        api_key_header="api-key",
        context_window_tokens=1_048_576,
        default_ocr_max_tokens=131072,
        default_translate_max_tokens=131072,
        default_summary_max_tokens=131072,
        default_translation_chunk_chars=120000,
        default_summary_chunk_chars=180000,
        thinking={"type": "enabled"},
        native_http=True,
        api_style="chat",
    ),
    "deepseek": ProviderProfile(
        core="deepseek",
        api_key_env="DEEPSEEK_API_KEY",
        base_url_env="DEEPSEEK_BASE_URL",
        model_env="DEEPSEEK_MODEL",
        default_base_url="https://api.deepseek.com",
        default_model="deepseek-v4-pro",
        token_param="max_tokens",
        default_vision_detail="high",
        supports_json_mode=True,
        supports_vision_detail=False,
        supports_image_input=False,
        api_key_header="authorization",
        context_window_tokens=1_000_000,
        default_ocr_max_tokens=12000,
        default_translate_max_tokens=384000,
        default_summary_max_tokens=384000,
        default_translation_chunk_chars=800000,
        default_summary_chunk_chars=850000,
        thinking={"type": "enabled"},
        reasoning_effort="high",
    ),
}


def load_env_file(path: Path) -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("export "):
            line = line[len("export ") :].strip()
        if "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def slugify(value: str, max_len: int = 80) -> str:
    value = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", value).strip(" ._")
    value = re.sub(r"_+", "_", value)
    return (value[:max_len] or "document").strip(" ._")


def now_stamp() -> str:
    return dt.datetime.now().strftime("%Y%m%d_%H%M%S")


def encode_image_data_url(image_path: Path) -> str:
    mime = mimetypes.guess_type(str(image_path))[0] or "image/jpeg"
    data = base64.b64encode(image_path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{data}"


def safe_json_loads(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.S)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass
        text_match = re.search(r'"text"\s*:\s*"(.*?)"\s*,\s*"notes"', cleaned, flags=re.S)
        if text_match:
            notes_match = re.search(r'"notes"\s*:\s*"(.*?)"\s*(?:\}|$)', cleaned, flags=re.S)
            return {
                "text": text_match.group(1).replace('\\"', '"').strip(),
                "notes": (notes_match.group(1).replace('\\"', '"').strip() if notes_match else ""),
            }
    return {"text": text.strip(), "notes": "Model did not return strict JSON."}


def retry_call(fn, attempts: int = 3, base_sleep: float = 2.0):
    last_exc: Exception | None = None
    for attempt in range(1, attempts + 1):
        try:
            return fn()
        except Exception as exc:  # noqa: BLE001 - surface provider errors after retries.
            last_exc = exc
            if isinstance(exc, (AuthenticationError, NativeAPIAuthenticationError)):
                break
            if attempt == attempts:
                break
            sleep_s = base_sleep * attempt
            print(f"[retry] API call failed on attempt {attempt}; retrying in {sleep_s:.1f}s: {type(exc).__name__}")
            time.sleep(sleep_s)
    raise RuntimeError(f"API call failed after {attempts} attempts: {last_exc}") from last_exc


class NativeAPIAuthenticationError(RuntimeError):
    pass


class LLMAgent:
    def __init__(
        self,
        *,
        profile: ProviderProfile,
        api_key: str,
        base_url: str,
        model: str,
        request_sleep: float = 0.2,
    ) -> None:
        self.client = None
        if not profile.native_http:
            client_kwargs: dict[str, Any] = {
                "api_key": api_key,
                "base_url": base_url,
            }
            self.client = OpenAI(**client_kwargs)
        self.profile = profile
        self.api_key = api_key
        self.base_url = base_url
        self.model = model
        self.request_sleep = request_sleep

    def chat_completion(self, kwargs: dict[str, Any]) -> str:
        if self.profile.native_http:
            if self.profile.api_style == "responses":
                return self.native_responses_completion(kwargs)
            return self.native_chat_completion(kwargs)
        if self.client is None:
            raise RuntimeError("OpenAI client is not initialized.")
        response = self.client.chat.completions.create(**kwargs)
        return response.choices[0].message.content or ""

    def native_chat_completion(self, kwargs: dict[str, Any]) -> str:
        payload: dict[str, Any] = {}
        for key, value in kwargs.items():
            if key == "extra_body":
                payload.update(value)
            else:
                payload[key] = value
        endpoint = self.base_url.rstrip("/") + "/chat/completions"
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers = {
            "Content-Type": "application/json",
            self.profile.api_key_header: self.api_key,
        }
        request = urllib.request.Request(endpoint, data=data, headers=headers, method="POST")
        try:
            with urllib.request.urlopen(request, timeout=180) as response:
                raw = response.read().decode("utf-8")
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            if exc.code in {401, 403}:
                raise NativeAPIAuthenticationError(f"HTTP {exc.code}: {detail}") from exc
            raise RuntimeError(f"HTTP {exc.code}: {detail}") from exc
        except urllib.error.URLError as exc:
            raise RuntimeError(f"Network error: {exc.reason}") from exc
        parsed = json.loads(raw)
        try:
            return parsed["choices"][0]["message"].get("content") or ""
        except (KeyError, IndexError, TypeError) as exc:
            raise RuntimeError(f"Unexpected API response: {raw[:1000]}") from exc

    def native_responses_completion(self, kwargs: dict[str, Any]) -> str:
        messages = kwargs.get("messages", [])
        instructions, response_input = self.to_responses_input(messages)
        payload: dict[str, Any] = {
            "model": kwargs["model"],
            "input": response_input,
        }
        if instructions:
            payload["instructions"] = instructions
        if self.profile.token_param in kwargs:
            payload["max_output_tokens"] = kwargs[self.profile.token_param]
        elif "max_tokens" in kwargs:
            payload["max_output_tokens"] = kwargs["max_tokens"]
        if "extra_body" in kwargs:
            payload.update(kwargs["extra_body"])
        endpoint = self.base_url.rstrip("/")
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}",
        }
        request = urllib.request.Request(endpoint, data=data, headers=headers, method="POST")
        try:
            with urllib.request.urlopen(request, timeout=180) as response:
                raw = response.read().decode("utf-8")
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            if exc.code in {401, 403}:
                raise NativeAPIAuthenticationError(f"HTTP {exc.code}: {detail}") from exc
            raise RuntimeError(f"HTTP {exc.code}: {detail}") from exc
        except urllib.error.URLError as exc:
            raise RuntimeError(f"Network error: {exc.reason}") from exc
        parsed = json.loads(raw)
        text = self.extract_responses_text(parsed)
        if text is None:
            raise RuntimeError(f"Unexpected Responses API response: {raw[:1000]}")
        return text

    @staticmethod
    def to_responses_input(messages: list[dict[str, Any]]) -> tuple[str, list[dict[str, Any]]]:
        instructions: list[str] = []
        converted_messages: list[dict[str, Any]] = []
        for message in messages:
            role = message.get("role", "user")
            content = message.get("content", "")
            if role == "system":
                if isinstance(content, str):
                    instructions.append(content)
                else:
                    instructions.append(json.dumps(content, ensure_ascii=False))
                continue
            converted_content: list[dict[str, Any]] = []
            if isinstance(content, str):
                converted_content.append({"type": "input_text", "text": content})
            elif isinstance(content, list):
                for part in content:
                    part_type = part.get("type")
                    if part_type == "text":
                        converted_content.append({"type": "input_text", "text": part.get("text", "")})
                    elif part_type == "image_url":
                        image_url = part.get("image_url", {})
                        converted_content.append(
                            {
                                "type": "input_image",
                                "image_url": image_url.get("url", ""),
                            }
                        )
            converted_messages.append({"role": role, "content": converted_content})
        return "\n\n".join(instructions), converted_messages

    @staticmethod
    def extract_responses_text(parsed: dict[str, Any]) -> str | None:
        if isinstance(parsed.get("output_text"), str):
            return parsed["output_text"]
        output = parsed.get("output")
        if isinstance(output, list):
            parts: list[str] = []
            for item in output:
                content = item.get("content") if isinstance(item, dict) else None
                if not isinstance(content, list):
                    continue
                for block in content:
                    if not isinstance(block, dict):
                        continue
                    if isinstance(block.get("text"), str):
                        parts.append(block["text"])
                    elif isinstance(block.get("output_text"), str):
                        parts.append(block["output_text"])
            if parts:
                return "\n".join(parts)
        return None

    def chat(
        self,
        messages: list[dict[str, Any]],
        max_tokens: int,
        *,
        json_mode: bool = False,
    ) -> str:
        use_json_mode = json_mode and self.profile.supports_json_mode

        def do_call(response_format: bool) -> str:
            kwargs: dict[str, Any] = {
                "model": self.model,
                "messages": messages,
                self.profile.token_param: max_tokens,
            }
            if self.profile.thinking:
                kwargs["extra_body"] = {"thinking": self.profile.thinking}
            if self.profile.reasoning_effort:
                kwargs["reasoning_effort"] = self.profile.reasoning_effort
            if response_format:
                kwargs["response_format"] = {"type": "json_object"}
            return self.chat_completion(kwargs)

        try:
            content = retry_call(lambda: do_call(use_json_mode))
        except RuntimeError:
            if not use_json_mode:
                raise
            content = retry_call(lambda: do_call(False))
        if self.request_sleep > 0:
            time.sleep(self.request_sleep)
        return content

    def image_url_payload(self, image_path: Path, vision_detail: str) -> dict[str, str]:
        payload = {"url": encode_image_data_url(image_path)}
        if self.profile.supports_vision_detail:
            payload["detail"] = vision_detail
        return payload

    def ocr_image(
        self,
        image_path: Path,
        *,
        page_label: str,
        embedded_text: str,
        vision_detail: str,
        max_tokens: int,
    ) -> tuple[str, str]:
        embedded_reference = embedded_text.strip()
        if len(embedded_reference) > 8000:
            embedded_reference = embedded_reference[:8000] + "\n...[embedded text truncated as reference]..."

        prompt = f"""
请对这页书籍/长文档图片进行严格 OCR，抽取所有可见文字。

要求：
1. 不要翻译，不要总结，不要改写。
2. 保留中文、日文、英文、数字、标点、专名、页码、脚注、表格文字和标题层级。
3. 按书籍阅读顺序输出：章节标题、正文段落、引文、脚注、页眉页脚、页码都要保留；页眉页脚可放在相邻位置，不要混入正文段落。
4. 日文竖排或从右到左排版时，按自然阅读顺序重排为横排文本；保留假名、汉字、外来语、专名和括号内原文。
5. 尽量保持原始段落和换行；表格、目录、索引、注释可转成 Markdown 表格或逐行文本。
6. 遇到模糊或无法辨认的字，用 [无法辨认] 标注，不要凭空补全。
7. 如果下方提供了 PDF 内嵌文字，只能把它作为校对参考；最终仍以图片中可见内容为准。
8. 如果本页没有正文，只输出可见的封面、版权页、目录页、空白页说明或页码信息。
9. 只输出 JSON，不要输出 Markdown 代码块。

JSON 格式：
{{
  "page": "{page_label}",
  "text": "完整抽取的原文",
  "notes": "无法辨认、版面异常或可能漏读的位置；没有则为空字符串"
}}

PDF 内嵌文字参考：
{embedded_reference if embedded_reference else "[无]"}
""".strip()

        messages = [
            {
                "role": "system",
                "content": "你是一个高精度书籍 OCR 引擎，任务是完整抽取中文、日文混排书页中的可见文字，并尽量保持书籍阅读顺序。",
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": self.image_url_payload(image_path, vision_detail),
                    },
                ],
            },
        ]
        raw = self.chat(messages, max_tokens=max_tokens, json_mode=True)
        data = safe_json_loads(raw)
        return str(data.get("text", "")).strip(), str(data.get("notes", "")).strip()

    def translate_page(
        self,
        text: str,
        *,
        page_label: str,
        target_language: str,
        max_tokens: int,
        chunk_chars: int,
    ) -> str:
        if not text.strip():
            return ""
        chunks = split_text(text, max(1000, chunk_chars))
        total_chunks = len(chunks)
        translations: list[str] = []
        for idx, chunk in enumerate(chunks, start=1):
            chunk_label = page_label if total_chunks == 1 else f"{page_label}，分段 {idx}/{total_chunks}"
            prompt = f"""
请将下面这段书籍 OCR 原文完整翻译为{target_language}。

要求：
1. 日文必须准确翻译；中文内容如果已经是简体中文，保留原意即可，不要扩写。
2. 不要总结，不要删减，不要解释。
3. 保留章节标题、段落、引文、脚注、页码、表格、目录层级和注释结构。
4. 人名、地名、书名、术语在同一文档中保持译名一致；不确定译名时保留原文并加 [存疑]。
5. 原文中的页码、章节号、脚注号、括号说明要原样保留，便于回查原页。
6. 对文学、哲学或学术类长文，优先保持语义准确和句间逻辑，不要为了流畅改写论证关系。
7. 无法确定的 OCR 内容保留原文并加 [存疑]。
8. 如果这是分段翻译，只翻译当前分段，不要补写其他分段内容，不要与其他页合并。
9. 只输出译文。

页码/分段：{chunk_label}

原文：
{chunk}
""".strip()
            messages = [
                {"role": "system", "content": "你是严谨的中日文书籍翻译助手，优先保证完整性、忠实度、术语一致性和页码可追溯性。"},
                {"role": "user", "content": prompt},
            ]
            translations.append(self.chat(messages, max_tokens=max_tokens).strip())
        return "\n\n".join(part for part in translations if part).strip()

    def summarize_chunk(self, text: str, *, chunk_label: str, max_tokens: int) -> str:
        prompt = f"""
请基于以下书籍/长文档内容做阶段性总结，输出简体中文。

要求：
1. 只依据给定内容，不要补充外部知识。
2. 这是整本书处理流程中的一个长分块，请提取能帮助最终整合的稳定信息。
3. 保留章节/小节线索、页码锚点、关键人物、术语、书名、引用来源、论点推进或叙事发展。
4. 区分“本段明确写到的内容”和“OCR/翻译不确定处”，不要推断未出现的信息。
5. 输出结构：
   - 范围与页码线索
   - 章节/段落结构
   - 主要内容
   - 关键人物/术语/译名
   - 论证或叙事推进
   - 需要复核之处

内容范围：{chunk_label}

文本：
{text}
""".strip()
        messages = [
            {"role": "system", "content": "你是书籍长文总结助手，擅长从分块文本中保留章节脉络、术语表和可回查页码线索。"},
            {"role": "user", "content": prompt},
        ]
        return self.chat(messages, max_tokens=max_tokens).strip()

    def final_summary(self, chunk_summaries: str, *, max_tokens: int) -> str:
        prompt = f"""
请整合下面的阶段性总结，生成整本书/长文档的最终总结，输出简体中文。

输出结构：
# 文档总结
## 一句话概括
## 全书结构
## 章节脉络
## 核心要点
## 重要人物/术语/译名
## 论证或叙事主线
## 可引用摘要
## OCR 与翻译复核提示

要求：
1. 不要编造阶段性总结中没有的信息。
2. 合并重复信息，保留章节顺序和跨页连续关系。
3. 对书籍类材料，重点说明作者/文本如何逐步展开观点或叙事。
4. 语言简洁，但信息要完整；不确定内容放入复核提示，不要写成事实。

阶段性总结：
{chunk_summaries}
""".strip()
        messages = [
            {"role": "system", "content": "你是书籍长文总结助手，负责把多个分块摘要整合成可阅读、可回查、结构清晰的最终总结。"},
            {"role": "user", "content": prompt},
        ]
        return self.chat(messages, max_tokens=max_tokens).strip()


def optimize_image(
    input_path: Path,
    output_path: Path,
    *,
    max_side: int,
    jpeg_quality: int,
) -> Path:
    with Image.open(input_path) as image:
        image = image.convert("RGB")
        width, height = image.size
        scale = min(1.0, max_side / max(width, height)) if max_side > 0 else 1.0
        if scale < 1.0:
            new_size = (max(1, int(width * scale)), max(1, int(height * scale)))
            image = image.resize(new_size, Image.Resampling.LANCZOS)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        image.save(output_path, "JPEG", quality=jpeg_quality, optimize=True)
    return output_path


def render_pdf_page_to_image(
    page: fitz.Page,
    output_path: Path,
    *,
    dpi: int,
    max_side: int,
    jpeg_quality: int,
) -> Path:
    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    mode = "RGB" if pix.n < 4 else "RGBA"
    image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
    if image.mode != "RGB":
        image = image.convert("RGB")
    width, height = image.size
    scale = min(1.0, max_side / max(width, height)) if max_side > 0 else 1.0
    if scale < 1.0:
        image = image.resize((max(1, int(width * scale)), max(1, int(height * scale))), Image.Resampling.LANCZOS)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path, "JPEG", quality=jpeg_quality, optimize=True)
    return output_path


def split_text(text: str, max_chars: int) -> list[str]:
    text = text.strip()
    if not text:
        return []
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    paragraphs = re.split(r"\n{2,}", text)
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        para_len = len(para)
        if current and current_len + para_len + 2 > max_chars:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
        if para_len > max_chars:
            for i in range(0, para_len, max_chars):
                part = para[i : i + max_chars]
                if current:
                    chunks.append("\n\n".join(current))
                    current = []
                    current_len = 0
                chunks.append(part)
        else:
            current.append(para)
            current_len += para_len + 2
    if current:
        chunks.append("\n\n".join(current))
    return chunks


def write_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


# ── 断点续传 (Checkpoint / Resume) 辅助函数 ──────────────────────────────────

def _checkpoint_dir(out_dir: Path) -> Path:
    """返回检查点目录路径，不存在则自动创建。"""
    d = out_dir / "_checkpoints"
    d.mkdir(parents=True, exist_ok=True)
    return d


def save_page_checkpoint(out_dir: Path, result: PageResult) -> None:
    """每完成一页 OCR 后，立即将结果写入 _checkpoints/page_XXXX.json。"""
    cp_dir = _checkpoint_dir(out_dir)
    path = cp_dir / f"page_{result.page_number:04d}.json"
    write_json(path, asdict(result))


def load_checkpoint_results(out_dir: Path) -> dict[int, PageResult]:
    """从 _checkpoints/ 加载所有已完成的页面结果，返回 {page_number: PageResult} 映射。"""
    cp_dir = _checkpoint_dir(out_dir)
    if not cp_dir.exists():
        return {}
    results: dict[int, PageResult] = {}
    for cp_file in sorted(cp_dir.glob("page_*.json")):
        try:
            data = json.loads(cp_file.read_text(encoding="utf-8"))
            pr = PageResult(**{k: v for k, v in data.items() if k in PageResult.__dataclass_fields__})
            results[pr.page_number] = pr
        except Exception as exc:
            print(f"[checkpoint] 跳过损坏的检查点 {cp_file.name}: {exc}")
    return results


def find_resume_output_dir(output_root: Path, pdf_name_stem: str) -> Path | None:
    """在 output_root 下查找与 pdf_name_stem 匹配的最新输出目录（用于 --resume）。

    搜索逻辑：遍历 output_root 下所有目录（含子目录中的 stamp 子目录），
    找到名称中包含 PDF 文件名的目录，按以下优先级排序：
      1. 同时有 _checkpoints 和 _page_images（最佳恢复源）
      2. 只有 _page_images 且图片数量多（可恢复 OCR）
      3. 只有 _checkpoints
      4. 都没有
    同级内按修改时间从新到旧排序。
    """
    if not output_root.exists():
        return None
    slug = slugify(pdf_name_stem)
    candidates: list[Path] = []

    # 1) 直接子目录
    for child in output_root.iterdir():
        if child.is_dir() and child.name.startswith(slug):
            candidates.append(child)

    # 2) 子目录中的 stamp 子目录（如 outputs/定本_完整处理/stamp_xxx/）
    for child in output_root.iterdir():
        if not child.is_dir():
            continue
        for grandchild in child.iterdir():
            if grandchild.is_dir() and grandchild.name.startswith(slug):
                candidates.append(grandchild)

    if not candidates:
        return None

    def _priority(p: Path) -> tuple:
        has_cp = (p / "_checkpoints").exists()
        cp_count = len(list((p / "_checkpoints").glob("page_*.json"))) if has_cp else 0
        has_img = (p / "_page_images").exists()
        img_count = len(list((p / "_page_images").glob("page_*.jpg"))) if has_img else 0
        # 优先级：有检查点 > 有图片且数量多 > 修改时间新
        return (cp_count, img_count, p.stat().st_mtime)

    candidates.sort(key=_priority, reverse=True)
    return candidates[0]


def recover_ocr_from_images(
    out_dir: Path,
    ocr_agent: LLMAgent,
    args: argparse.Namespace,
) -> int:
    """从已有的 _page_images/page_XXXX.jpg 恢复 OCR 检查点。

    用于处理之前运行中断但未保存检查点的场景。
    图片文件名格式为 page_XXXX.jpg，XXXX 即页码。
    返回恢复的页面数。
    """
    cache_dir = out_dir / "_page_images"
    if not cache_dir.exists():
        return 0

    existing = load_checkpoint_results(out_dir)
    recovered = 0
    images = sorted(cache_dir.glob("page_*.jpg"))
    total = len(images)

    for img_path in images:
        # 从文件名提取页码：page_0157.jpg → 157
        match = re.match(r"page_(\d+)\.jpg", img_path.name)
        if not match:
            continue
        page_number = int(match.group(1))
        if page_number in existing:
            continue  # 已有检查点，跳过

        print(f"[recover] 正在恢复第 {page_number} 页 OCR ({total} 张图片)")
        # 使用嵌入文本作为参考（如果有 PDF 的话可以通过 fitz 获取）
        text, notes = ocr_agent.ocr_image(
            img_path,
            page_label=f"第 {page_number} 页",
            embedded_text="",  # 无法直接获取嵌入文本，留空
            vision_detail=args.vision_detail,
            max_tokens=args.ocr_max_tokens,
        )
        result = PageResult(
            source=str(img_path),
            kind="pdf_page",
            page_number=page_number,
            ocr_text=text,
            notes=notes,
            image_path=str(img_path),
        )
        save_page_checkpoint(out_dir, result)
        recovered += 1
    return recovered


def _page_has_translation(result: PageResult) -> bool:
    """判断该页翻译是否已完成（非空即视为已完成）。"""
    return bool(result.translation.strip())


def resolve_profile(core_value: str | None) -> ProviderProfile:
    core = (core_value or "seed").strip().lower()
    if core not in PROVIDER_PROFILES:
        choices = ", ".join(PROVIDER_PROFILES)
        raise ValueError(f"Unsupported LLM core: {core_value}. Choose one of: {choices}.")
    return PROVIDER_PROFILES[core]


def resolve_agent_connection(
    *,
    profile: ProviderProfile,
    api_key_override: str | None = None,
    base_url_override: str | None = None,
    model_override: str | None = None,
    vision_detail_override: str | None = None,
) -> tuple[str, str, str, str]:
    api_key = api_key_override or os.getenv(profile.api_key_env, "")
    if profile.core == "seed" and not api_key:
        api_key = os.getenv("OPENAI_API_KEY", "")
    base_url = base_url_override or os.getenv(profile.base_url_env, profile.default_base_url)
    model = model_override or os.getenv(profile.model_env, profile.default_model)
    vision_detail = vision_detail_override or os.getenv(
        f"{profile.core.upper()}_VISION_DETAIL",
        profile.default_vision_detail,
    )
    return api_key, base_url, model, vision_detail


def resolve_provider_settings(args: argparse.Namespace) -> tuple[ProviderProfile, str, str, str, str]:
    profile = resolve_profile(args.llm_core)
    api_key, base_url, model, vision_detail = resolve_agent_connection(
        profile=profile,
        api_key_override=args.api_key,
        base_url_override=args.api_base,
        model_override=args.model,
        vision_detail_override=args.vision_detail,
    )
    args.translate_max_tokens = args.translate_max_tokens or int(
        os.getenv(f"{profile.core.upper()}_TRANSLATE_MAX_TOKENS", profile.default_translate_max_tokens)
    )
    args.summary_max_tokens = args.summary_max_tokens or int(
        os.getenv(f"{profile.core.upper()}_SUMMARY_MAX_TOKENS", profile.default_summary_max_tokens)
    )
    args.translation_chunk_chars = args.translation_chunk_chars or int(
        os.getenv(f"{profile.core.upper()}_TRANSLATION_CHUNK_CHARS", profile.default_translation_chunk_chars)
    )
    args.summary_chunk_chars = args.summary_chunk_chars or int(
        os.getenv(f"{profile.core.upper()}_SUMMARY_CHUNK_CHARS", profile.default_summary_chunk_chars)
    )
    return profile, api_key, base_url, model, vision_detail


def resolve_ocr_provider_settings(
    args: argparse.Namespace,
    text_profile: ProviderProfile,
) -> tuple[ProviderProfile, str, str, str, str]:
    ocr_core = args.ocr_llm_core or os.getenv("OCR_LLM_CORE")
    if not ocr_core:
        ocr_core = text_profile.core if text_profile.supports_image_input else "mimo"
    profile = resolve_profile(ocr_core)
    if not profile.supports_image_input:
        raise ValueError(f"OCR LLM core '{profile.core}' does not support image input. Choose seed or mimo for OCR.")
    api_key, base_url, model, vision_detail = resolve_agent_connection(
        profile=profile,
        vision_detail_override=args.vision_detail,
    )
    args.ocr_max_tokens = args.ocr_max_tokens or int(
        os.getenv(f"{profile.core.upper()}_OCR_MAX_TOKENS", profile.default_ocr_max_tokens)
    )
    return profile, api_key, base_url, model, vision_detail


def page_markdown(results: Iterable[PageResult], field: str) -> str:
    parts: list[str] = []
    for result in results:
        value = getattr(result, field).strip()
        if not value:
            continue
        parts.append(f"## 第 {result.page_number} 页\n\n{value}")
        if result.notes and field == "ocr_text":
            parts.append(f"> OCR 备注：{result.notes}")
    return "\n\n".join(parts).strip() + "\n"


def set_docx_font(document: Document, font_name: str = "Microsoft YaHei") -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = font_name
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_markdownish_to_docx(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\s*[-*]\s+", line):
            document.add_paragraph(re.sub(r"^\s*[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\s*\d+[.、]\s+", line):
            document.add_paragraph(line, style="List Number")
        else:
            document.add_paragraph(line)


def build_docx(
    output_path: Path,
    *,
    title: str,
    metadata: dict[str, Any],
    summary: str,
    translation: str,
    extracted: str,
) -> None:
    document = Document()
    set_docx_font(document)
    document.add_heading(title, level=0)
    document.add_heading("处理信息", level=1)
    for key, value in metadata.items():
        document.add_paragraph(f"{key}: {value}")

    document.add_page_break()
    document.add_heading("内容总结", level=1)
    add_markdownish_to_docx(document, summary or "未生成总结。")

    document.add_page_break()
    document.add_heading("中文翻译", level=1)
    add_markdownish_to_docx(document, translation or "未生成翻译。")

    document.add_page_break()
    document.add_heading("原文抽取", level=1)
    add_markdownish_to_docx(document, extracted or "未生成原文抽取。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def register_pdf_font() -> str:
    candidates = [
        ("MicrosoftYaHei", Path(r"C:\Windows\Fonts\msyh.ttc")),
        ("SimSun", Path(r"C:\Windows\Fonts\simsun.ttc")),
        ("MSGothic", Path(r"C:\Windows\Fonts\msgothic.ttc")),
        ("YuGothic", Path(r"C:\Windows\Fonts\YuGothR.ttc")),
    ]
    for font_name, font_path in candidates:
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            return font_name
        except Exception:
            continue
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def make_pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CJKTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=TA_LEFT,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "CJKHeading1",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=20,
            spaceBefore=10,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "CJKHeading2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=12,
            leading=18,
            spaceBefore=8,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "CJKBody",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=9.5,
            leading=14,
            textColor=colors.black,
            wordWrap="CJK",
        ),
        "meta": ParagraphStyle(
            "CJKMeta",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor("#444444"),
            wordWrap="CJK",
        ),
    }


def add_text_to_pdf_story(story: list[Any], text: str, styles: dict[str, ParagraphStyle]) -> None:
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        escaped = html.escape(line)
        if line.startswith("# "):
            story.append(Paragraph(escaped[2:].strip(), styles["h1"]))
        elif line.startswith("## "):
            story.append(Paragraph(escaped[3:].strip(), styles["h2"]))
        elif line.startswith("### "):
            story.append(Paragraph(escaped[4:].strip(), styles["h2"]))
        else:
            # Very long OCR lines can make ReportLab slow; split without changing content.
            for start in range(0, len(escaped), 1800):
                story.append(Paragraph(escaped[start : start + 1800], styles["body"]))


def build_pdf(
    output_path: Path,
    *,
    title: str,
    metadata: dict[str, Any],
    summary: str,
    translation: str,
    extracted: str,
) -> None:
    font_name = register_pdf_font()
    styles = make_pdf_styles(font_name)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    story: list[Any] = [Paragraph(html.escape(title), styles["title"]), Spacer(1, 8)]
    story.append(Paragraph("处理信息", styles["h1"]))
    for key, value in metadata.items():
        story.append(Paragraph(html.escape(f"{key}: {value}"), styles["meta"]))

    story.extend([PageBreak(), Paragraph("内容总结", styles["h1"])])
    add_text_to_pdf_story(story, summary or "未生成总结。", styles)

    story.extend([PageBreak(), Paragraph("中文翻译", styles["h1"])])
    add_text_to_pdf_story(story, translation or "未生成翻译。", styles)

    story.extend([PageBreak(), Paragraph("原文抽取", styles["h1"])])
    add_text_to_pdf_story(story, extracted or "未生成原文抽取。", styles)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def process_pdf(path: Path, out_dir: Path, args: argparse.Namespace, ocr_agent: LLMAgent | None) -> list[PageResult]:
    results: list[PageResult] = []
    cache_dir = out_dir / "_page_images"

    # ── 断点续传：加载已有检查点 ──
    existing = load_checkpoint_results(out_dir) if args.resume else {}
    # ── OCR 恢复：从已有页面图片重建检查点 ──
    if getattr(args, "ocr_recover", False) and ocr_agent is not None and not args.skip_ocr:
        recover_count = recover_ocr_from_images(out_dir, ocr_agent, args)
        if recover_count:
            print(f"[recover] 从页面图片恢复了 {recover_count} 页 OCR 检查点")
        # 重新加载检查点（包含新恢复的）
        existing = load_checkpoint_results(out_dir)
    if existing:
        skipped = len(existing)
        max_done = max(existing)
        print(f"[resume] 已加载 {skipped} 个已完成页面的检查点（最大页码: {max_done}）")
    else:
        max_done = 0

    with fitz.open(path) as document:
        total_pages = document.page_count
        # 续传时从已完成页的下一页开始；否则使用 start_page
        if existing and args.resume:
            start = max_done + 1
        else:
            start = max(1, args.start_page)
        end = total_pages if args.max_pages is None else min(total_pages, start + args.max_pages - 1)

        for page_number in range(start, end + 1):
            # ── 断点续传：跳过已有检查点的页面 ──
            if page_number in existing:
                results.append(existing[page_number])
                print(f"[resume] 跳过已完成的第 {page_number} 页")
                continue

            page = document.load_page(page_number - 1)
            embedded_text = page.get_text("text").strip()
            result = PageResult(
                source=str(path),
                kind="pdf_page",
                page_number=page_number,
                embedded_text=embedded_text,
            )
            print(f"[ocr] PDF page {page_number}/{total_pages}")
            if not args.skip_ocr:
                if ocr_agent is None:
                    raise ValueError(f"OCR requires an API key. Set {args.ocr_provider_profile.api_key_env} or pass --skip-ocr.")
                image_path = cache_dir / f"page_{page_number:04d}.jpg"
                render_pdf_page_to_image(
                    page,
                    image_path,
                    dpi=args.dpi,
                    max_side=args.max_image_side,
                    jpeg_quality=args.jpeg_quality,
                )
                text, notes = ocr_agent.ocr_image(
                    image_path,
                    page_label=f"第 {page_number} 页",
                    embedded_text=embedded_text,
                    vision_detail=args.vision_detail,
                    max_tokens=args.ocr_max_tokens,
                )
                result.ocr_text = text
                result.notes = notes
                result.image_path = str(image_path)
                # ── 断点续传：保存该页检查点 ──
                save_page_checkpoint(out_dir, result)
            results.append(result)

    # 如果续传并重新处理了尾部页面，把之前已跳过的页面也合并进来
    if existing:
        all_results = {r.page_number: r for r in results}
        for pg, pr in existing.items():
            if pg not in all_results:
                all_results[pg] = pr
        results = [all_results[pg] for pg in sorted(all_results)]

    return results


def process_image(path: Path, out_dir: Path, args: argparse.Namespace, ocr_agent: LLMAgent | None, page_number: int) -> PageResult:
    result = PageResult(source=str(path), kind="image", page_number=page_number)
    print(f"[ocr] Image {path.name}")
    if not args.skip_ocr:
        if ocr_agent is None:
            raise ValueError(f"OCR requires an API key. Set {args.ocr_provider_profile.api_key_env} or pass --skip-ocr.")
        cache_dir = out_dir / "_page_images"
        image_path = cache_dir / f"image_{page_number:04d}.jpg"
        optimize_image(path, image_path, max_side=args.max_image_side, jpeg_quality=args.jpeg_quality)
        text, notes = ocr_agent.ocr_image(
            image_path,
            page_label=f"图片 {page_number}",
            embedded_text="",
            vision_detail=args.vision_detail,
            max_tokens=args.ocr_max_tokens,
        )
        result.ocr_text = text
        result.notes = notes
        result.image_path = str(image_path)
    return result


def translate_results(results: list[PageResult], args: argparse.Namespace, text_agent: LLMAgent | None, out_dir: Path | None = None) -> None:
    if args.skip_translation:
        return
    if text_agent is None:
        raise ValueError(f"Translation requires an API key. Set {args.provider_profile.api_key_env} or pass --skip-translation.")
    total = len(results)
    for idx, result in enumerate(results, start=1):
        # ── 断点续传：跳过已有翻译的页面 ──
        if args.resume and _page_has_translation(result):
            print(f"[resume] 跳过已翻译的第 {result.page_number} 页")
            continue
        print(f"[translate] page {result.page_number} ({idx}/{total})")
        result.translation = text_agent.translate_page(
            result.source_text,
            page_label=f"第 {result.page_number} 页",
            target_language=args.target_language,
            max_tokens=args.translate_max_tokens,
            chunk_chars=args.translation_chunk_chars,
        )
        # 每翻译一页保存一次检查点（如果 out_dir 可用）
        if out_dir is not None:
            save_page_checkpoint(out_dir, result)


def summarize_results(results: list[PageResult], args: argparse.Namespace, text_agent: LLMAgent | None) -> str:
    if args.skip_summary:
        return ""
    if text_agent is None:
        raise ValueError(f"Summary requires an API key. Set {args.provider_profile.api_key_env} or pass --skip-summary.")
    source = page_markdown(results, "translation").strip() or page_markdown(results, "ocr_text").strip()
    if not source:
        return ""
    chunks = split_text(source, args.summary_chunk_chars)
    chunk_summaries: list[str] = []
    for idx, chunk in enumerate(chunks, start=1):
        print(f"[summary] chunk {idx}/{len(chunks)}")
        chunk_summaries.append(
            text_agent.summarize_chunk(
                chunk,
                chunk_label=f"分块 {idx}/{len(chunks)}",
                max_tokens=args.summary_max_tokens,
            )
        )
    joined = "\n\n".join(f"## 分块 {i}\n{summary}" for i, summary in enumerate(chunk_summaries, start=1))
    print("[summary] final")
    return text_agent.final_summary(joined, max_tokens=args.summary_max_tokens)


def process_input(
    path: Path,
    args: argparse.Namespace,
    *,
    ocr_agent: LLMAgent | None,
    text_agent: LLMAgent | None,
) -> Path:
    if not path.exists():
        raise FileNotFoundError(path)
    output_root = Path(args.output_dir)

    # ── 断点续传：定位已有输出目录 ──
    out_dir: Path
    if getattr(args, "resume", False) or getattr(args, "resume_from", None):
        if getattr(args, "resume_from", None):
            # 用户显式指定了续传目录
            resume_path = Path(args.resume_from)
            if not resume_path.is_absolute():
                resume_path = output_root / resume_path
            out_dir = resume_path
        else:
            out_dir = find_resume_output_dir(output_root, path.stem)
        if out_dir is None or not out_dir.exists():
            msg = f"找不到可续传的输出目录。请检查 --output-dir ({output_root}) 或使用 --resume-from 指定目录。"
            raise FileNotFoundError(msg)
        print(f"[resume] 使用已有输出目录: {out_dir.name}")
    else:
        out_dir = output_root / f"{slugify(path.stem)}_{now_stamp()}"
        out_dir.mkdir(parents=True, exist_ok=True)

    suffix = path.suffix.lower()
    if suffix in PDF_EXTENSIONS:
        results = process_pdf(path, out_dir, args, ocr_agent)
    elif suffix in IMAGE_EXTENSIONS:
        results = [process_image(path, out_dir, args, ocr_agent, 1)]
    else:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    translate_results(results, args, text_agent, out_dir=out_dir)
    summary = summarize_results(results, args, text_agent)

    extracted_md = page_markdown(results, "ocr_text").strip()
    if not extracted_md:
        extracted_md = page_markdown(results, "embedded_text").strip()
    translation_md = page_markdown(results, "translation").strip()

    metadata = {
        "input": str(path),
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "llm_core": args.provider_profile.core,
        "provider_model": args.provider_model,
        "provider_base_url": args.provider_base_url,
        "ocr_llm_core": args.ocr_provider_profile.core,
        "ocr_provider_model": args.ocr_provider_model,
        "ocr_provider_base_url": args.ocr_provider_base_url,
        "context_window_tokens": args.provider_profile.context_window_tokens,
        "thinking": args.provider_profile.thinking or {"type": "default"},
        "reasoning_effort": args.provider_profile.reasoning_effort,
        "ocr_context_window_tokens": args.ocr_provider_profile.context_window_tokens,
        "ocr_thinking": args.ocr_provider_profile.thinking or {"type": "default"},
        "ocr_max_tokens": args.ocr_max_tokens,
        "translate_max_tokens": args.translate_max_tokens,
        "summary_max_tokens": args.summary_max_tokens,
        "translation_chunk_chars": args.translation_chunk_chars,
        "summary_chunk_chars": args.summary_chunk_chars,
        "vision_detail": args.vision_detail,
        "items": len(results),
        "target_language": args.target_language,
    }

    write_json(out_dir / "metadata.json", metadata)
    write_json(out_dir / "extracted_pages.json", [asdict(result) for result in results])
    (out_dir / "extracted_text.md").write_text(extracted_md + "\n", encoding="utf-8")
    (out_dir / "translation.md").write_text(translation_md + "\n", encoding="utf-8")
    (out_dir / "summary.md").write_text(summary + "\n", encoding="utf-8")

    title = f"{path.stem} - OCR 翻译总结"
    build_docx(
        out_dir / "summary.docx",
        title=title,
        metadata=metadata,
        summary=summary,
        translation=translation_md,
        extracted=extracted_md,
    )
    build_pdf(
        out_dir / "summary.pdf",
        title=title,
        metadata=metadata,
        summary=summary,
        translation=translation_md,
        extracted=extracted_md,
    )

    if not args.keep_page_images:
        cache_dir = out_dir / "_page_images"
        if cache_dir.exists() and cache_dir.is_dir():
            shutil.rmtree(cache_dir)

    print(f"[done] {out_dir}")
    return out_dir


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Extract OCR text from PDFs/images with a switchable LLM core, translate, summarize, and export DOCX/PDF.",
    )
    parser.add_argument("inputs", nargs="+", help="PDF or image files to process.")
    parser.add_argument("-o", "--output-dir", default="outputs", help="Output directory.")
    parser.add_argument("--target-language", default="简体中文", help="Translation target language.")
    parser.add_argument(
        "--llm-core",
        default=os.getenv("LLM_CORE", "seed"),
        choices=["seed", "mimo", "deepseek"],
        help="Text LLM core used for translation and summary.",
    )
    parser.add_argument(
        "--ocr-llm-core",
        default=os.getenv("OCR_LLM_CORE"),
        choices=["seed", "mimo"],
        help="Multimodal LLM core used for PDF/image OCR. Defaults to text core if it supports images; otherwise mimo.",
    )
    parser.add_argument("--api-key", default=None, help="Advanced override for the selected provider API key.")
    parser.add_argument("--api-base", default=None, help="Advanced override for the selected provider base URL.")
    parser.add_argument("--model", default=None, help="Advanced override for the selected provider model ID.")
    parser.add_argument(
        "--vision-detail",
        default=None,
        choices=["low", "high", "xhigh"],
        help="Vision detail for providers that support it. Seed defaults to xhigh.",
    )
    parser.add_argument("--start-page", type=int, default=1, help="1-based PDF start page.")
    parser.add_argument("--max-pages", type=int, default=None, help="Maximum PDF pages to process.")
    parser.add_argument("--dpi", type=int, default=220, help="PDF render DPI before OCR.")
    parser.add_argument("--max-image-side", type=int, default=3200, help="Resize OCR image max side.")
    parser.add_argument("--jpeg-quality", type=int, default=92, help="JPEG quality for OCR images.")
    parser.add_argument("--ocr-max-tokens", type=int, default=None, help="Override provider OCR output token budget.")
    parser.add_argument("--translate-max-tokens", type=int, default=None, help="Override provider translation token budget.")
    parser.add_argument("--translation-chunk-chars", type=int, default=None, help="Override translation chunk size.")
    parser.add_argument("--summary-max-tokens", type=int, default=None, help="Override provider summary token budget.")
    parser.add_argument("--summary-chunk-chars", type=int, default=None, help="Override summary chunk size.")
    parser.add_argument("--request-sleep", type=float, default=0.2, help="Sleep seconds after each API request.")
    parser.add_argument("--skip-ocr", action="store_true", help="Use PDF embedded text only; do not call vision OCR.")
    parser.add_argument("--skip-translation", action="store_true")
    parser.add_argument("--skip-summary", action="store_true")
    parser.add_argument("--keep-page-images", action="store_true", help="Keep rendered page images for audit.")
    parser.add_argument(
        "--resume", action="store_true",
        help="Resume from the latest output directory. Loads existing OCR/translation checkpoints and skips completed pages.",
    )
    parser.add_argument(
        "--resume-from", default=None,
        help="Resume from a specific output directory (full path or directory name under --output-dir).",
    )
    parser.add_argument(
        "--ocr-recover", action="store_true",
        help="Re-run OCR on existing page images in _page_images/ to rebuild checkpoints (for recovering from interrupted runs that lacked checkpoints).",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    load_env_file(Path(__file__).with_name(".env"))
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        profile, api_key, base_url, model, vision_detail = resolve_provider_settings(args)
        if args.skip_ocr:
            requested_ocr_core = args.ocr_llm_core or os.getenv("OCR_LLM_CORE")
            ocr_profile = resolve_profile(requested_ocr_core) if requested_ocr_core else (
                profile if profile.supports_image_input else PROVIDER_PROFILES["mimo"]
            )
            if not ocr_profile.supports_image_input:
                ocr_profile = PROVIDER_PROFILES["mimo"]
            ocr_api_key = ""
            _, ocr_base_url, ocr_model, ocr_vision_detail = resolve_agent_connection(
                profile=ocr_profile,
                vision_detail_override=args.vision_detail,
            )
            args.ocr_max_tokens = args.ocr_max_tokens or int(
                os.getenv(f"{ocr_profile.core.upper()}_OCR_MAX_TOKENS", ocr_profile.default_ocr_max_tokens)
            )
        else:
            ocr_profile, ocr_api_key, ocr_base_url, ocr_model, ocr_vision_detail = resolve_ocr_provider_settings(args, profile)
    except ValueError as exc:
        parser.error(str(exc))
    args.provider_profile = profile
    args.provider_base_url = base_url
    args.provider_model = model
    args.ocr_provider_profile = ocr_profile
    args.ocr_provider_base_url = ocr_base_url
    args.ocr_provider_model = ocr_model
    args.vision_detail = ocr_vision_detail

    text_agent: LLMAgent | None = None
    if not (args.skip_translation and args.skip_summary):
        if not api_key:
            parser.error(f"Missing API key. Set {profile.api_key_env} in .env or pass --api-key.")
        text_agent = LLMAgent(
            profile=profile,
            api_key=api_key,
            base_url=base_url,
            model=model,
            request_sleep=args.request_sleep,
        )

    ocr_agent: LLMAgent | None = None
    if not args.skip_ocr:
        if not ocr_api_key:
            parser.error(f"Missing OCR API key. Set {ocr_profile.api_key_env} in .env or pass --skip-ocr.")
        ocr_agent = LLMAgent(
            profile=ocr_profile,
            api_key=ocr_api_key,
            base_url=ocr_base_url,
            model=ocr_model,
            request_sleep=args.request_sleep,
        )

    failures: list[str] = []
    for input_value in args.inputs:
        path = Path(input_value).expanduser().resolve()
        try:
            process_input(path, args, ocr_agent=ocr_agent, text_agent=text_agent)
        except Exception as exc:  # noqa: BLE001 - continue processing later inputs.
            failures.append(f"{path}: {type(exc).__name__}: {exc}")
            print(f"[error] {failures[-1]}", file=sys.stderr)

    if failures:
        print("\nFailures:", file=sys.stderr)
        for failure in failures:
            print(f"- {failure}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
