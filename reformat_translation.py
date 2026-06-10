# -*- coding: utf-8 -*-
"""
将 translation.md 按自然段和文章标题重新整合为 Word 文档。
去除页面标记和页码，保留文章标题层级结构。
"""
import re
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 确保 UTF-8 输出
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

INPUT_FILE = Path("outputs/定本_完整处理/定本 柄谷行人文学论集 (柄谷行人 著，陈言 译) (z-library.sk, 1lib.sk, z-lib.sk)_20260609_155715/translation.md")
OUTPUT_FILE = Path("outputs/定本_完整处理/定本 柄谷行人文学论集 (柄谷行人 著，陈言 译) (z-library.sk, 1lib.sk, z-lib.sk)_20260609_155715/translation_formatted.docx")


def is_page_marker(line: str) -> bool:
    """判断是否为页面标记行（如 ## 第 1 页）"""
    return bool(re.match(r'^## 第 \d+ 页', line))


def is_page_number(line: str) -> bool:
    """判断是否为孤立的页码行（如 '第 6 页' 或 '序文 | 001'）"""
    line = line.strip()
    # 纯页码行：第 X 页
    if re.match(r'^第 \d+ 页$', line):
        return True
    # 页眉页脚格式：章节名 | 页码
    if re.match(r'^.+\| \d+$', line):
        return True
    # 纯数字页码（独立一行）
    if re.match(r'^\d{1,4}$', line):
        return True
    return False


def is_noise_line(line: str) -> bool:
    """判断是否为噪声行（ISBN、定价、出版社信息、封面信息、页眉页脚等）"""
    line = line.strip()
    noise_patterns = [
        r'^ISBN \d',
        r'^定价[：:]\s*\d',
        r'^微信扫描',
        r'^上架建议',
        r'^9 \d{10,}',
        r'^新浪微博',
        r'^微\s*信[：:]',
        r'^淘宝店铺',
        r'^本社常年法律顾问',
        r'^凡有印装质量问题',
        r'^责任编辑',
        r'^责任印制',
        r'^出版发行',
        r'^地\s+址[：:]',
        r'^电\s+话[：:]',
        r'^传\s+真[：:]',
        r'^经\s+销[：:]',
        r'^印\s+刷[：:]',
        r'^开\s+本[：:]',
        r'^字\s+数[：:]',
        r'^印\s+张[：:]',
        r'^版\s+次[：:]',
        r'^印\s+次[：:]',
        r'^ICCTP',
        r'^Central Compilation',
        r'^著$',
        r'^译$',
        r'^陈言\s*译$',
        r'^柄谷行人\s*著$',
        r'^中央编译出版社',
        r'^定本$',
        r'^柄谷行人文学论集$',
        r'^柄谷行人$',
        r'^定价[：:]',
        r'^图书在版编目',
        r'^著作权登记号',
        r'^TEIHON',
        r'^by Kojin',
        r'^Originally published',
        r'^This simplified',
        r'^by arrangement',
        r'^责任编辑[：:]',
        r'^\d+页$',
        r'^第 \d+ 页$',
        r'^\d{3,4}\s*\|\s*定本',  # 页眉页脚格式
        r'^\d{3,4}\s*\|\s*序文',
        r'^\d{3,4}\s*\|\s*漱石',
        r'^\d{3,4}\s*\|\s*麦克白',
        r'^\d{3,4}\s*\|\s*坂口',
        r'^\d{3,4}\s*\|\s*梦的',
        r'^\d{3,4}\s*\|\s*中上',
        r'^\d{3,4}\s*\|\s*文学',
        r'^\d{3,4}\s*\|\s*马克思',
        r'^\d{3,4}\s*\|\s*柳田',
        r'^\d{3,4}\s*\|\s*鲁迅',
        r'^\d{3,4}\s*\|\s*芥川',
        r'^\d{3,4}\s*\|\s*三岛',
        r'^\d{3,4}\s*\|\s*二叶亭',
        r'^\d{3,4}\s*\|\s*近代',
        r'^\d{3,4}\s*\|\s*翻译',
        r'^\d{3,4}\s*\|\s*森鸥外',
        r'^\d{3,4}\s*\|\s*亚历山大',
        r'^9 \d{6,}',  # ISBN 条形码数字
        r'^定本柄谷行人文文学论集 \[存疑\]',  # 版权页标题
        r'^作者简介$',
        r'^译者简介$',
        r'^序文 \| \d+$',  # 页眉页脚
        r'^漱石试论 \| \d+$',
        r'^麦克白 \| \d+$',
        r'^坂口安吾 \| \d+$',
        r'^梦的世界 \| \d+$',
        r'^中上健次 \| \d+$',
        r'^文学的衰灭 \| \d+$',
        r'^马克思 \| \d+$',
        r'^柳田国男 \| \d+$',
        r'^鲁迅 \| \d+$',
        r'^芥川 \| \d+$',
        r'^三岛 \| \d+$',
        r'^二叶亭 \| \d+$',
        r'^近代文学 \| \d+$',
        r'^翻译的方法 \| \d+$',
        r'^森鸥外 \| \d+$',
        r'^亚历山大 \| \d+$',
    ]
    for pattern in noise_patterns:
        if re.match(pattern, line):
            return True
    return False


def detect_title(line: str) -> int | None:
    """
    检测文章标题层级。
    返回: 1=一级标题, 2=二级标题, 3=三级标题, None=非标题
    """
    line = line.strip()
    if not line:
        return None

    # 一级标题：书名、大章节（如 "定本 柄谷行人文学论集"、"序 文"、"第一部"、"第二部"）
    level1_patterns = [
        r'^定本\s*柄谷行人文学论集',
        r'^序\s*文$',
        r'^第一部$',
        r'^第二部$',
    ]
    for p in level1_patterns:
        if re.match(p, line):
            return 1

    # 二级标题：具体文章名（精确匹配，避免对话中的误判）
    # 特征：独立成行的论文/章节标题
    level2_patterns = [
        r'^漱石试论——意识与自然$',
        r'^夏目漱石论$',
        r'^麦克白论$',
        r'^森鸥外历史小说论$',
        r'^坂口安吾，其可能性的中心$',
        r'^梦的世界——岛尾敏雄$',
        r'^中上健次与福克纳$',
        r'^二叶亭四迷的翻译方法$',
        r'^近代文学的终焉$',
        r'^文学的衰灭$',
        r'^马克思，其可能性的中心$',
        r'^柳田国男试论$',
        r'^鲁迅的复古与革新$',
        r'^芥川龙之介论$',
        r'^三岛由纪夫论$',
        r'^亚历山大四重奏的辩证法$',
    ]
    for p in level2_patterns:
        if re.match(p, line):
            return 2

    # 三级标题：小节标题
    level3_patterns = [
        r'^一[、．.]',
        r'^二[、．.]',
        r'^三[、．.]',
        r'^四[、．.]',
        r'^五[、．.]',
        r'^\d+[、．.]',
    ]
    for p in level3_patterns:
        if re.match(p, line):
            return 3

    return None


def clean_content(lines: list[str]) -> list[str]:
    """清理内容：去除页面标记、页码、噪声行、封面信息"""
    cleaned = []
    in_copyright = False  # 跟踪是否在版权页区域
    skip_first_pages = True  # 跳过前几页（封面、版权页等）

    for line in lines:
        stripped = line.strip()

        # 跳过页面标记
        if is_page_marker(line):
            # 检查页面编号
            page_match = re.match(r'^## 第 (\d+) 页', line)
            if page_match:
                page_num = int(page_match.group(1))
                # 前5页跳过
                if page_num <= 5:
                    continue
                # 版权页区域
                if in_copyright:
                    if page_num <= 5:
                        continue
                    else:
                        in_copyright = False
                skip_first_pages = False
            continue

        # 检测版权页开始
        if re.match(r'^定本\s*柄谷行人文学论集\s*\[存疑\]', stripped):
            in_copyright = True
            continue
        if re.match(r'^图书在版编目', stripped):
            in_copyright = True
            continue

        # 跳过版权页内容
        if in_copyright:
            if re.match(r'^序\s*文$', stripped):
                in_copyright = False
            else:
                continue

        # 跳过前几页的残留内容
        if skip_first_pages:
            continue

        # 跳过页码
        if is_page_number(line):
            continue

        # 跳过噪声行
        if is_noise_line(line):
            continue

        cleaned.append(line)
    return cleaned


def merge_paragraphs(lines: list[str]) -> list[tuple[str, int]]:
    """
    合并自然段落，返回 [(内容, 标题层级)] 列表。
    标题层级: 0=正文, 1=一级标题, 2=二级标题, 3=三级标题
    """
    result = []
    current_para = []
    seen_titles = set()  # 跟踪已出现的标题，避免重复

    for line in lines:
        stripped = line.strip()

        # 空行：段落分隔
        if not stripped:
            if current_para:
                result.append(("\n".join(current_para), 0))
                current_para = []
            continue

        # 检测标题
        title_level = detect_title(stripped)
        if title_level is not None:
            # 先保存当前段落
            if current_para:
                result.append(("\n".join(current_para), 0))
                current_para = []
            # 检查标题是否已出现过（避免目录和正文重复）
            if stripped not in seen_titles:
                seen_titles.add(stripped)
                result.append((stripped, title_level))
            continue

        # 脚注标记（如 ①、② 等）：保留在当前段落
        if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', stripped):
            if current_para:
                current_para.append(stripped)
            else:
                current_para.append(stripped)
            continue

        # 普通文本行：追加到当前段落
        current_para.append(stripped)

    # 保存最后一个段落
    if current_para:
        result.append(("\n".join(current_para), 0))

    return result


def set_docx_font(document: Document, font_name: str = "Microsoft YaHei") -> None:
    """设置文档默认字体"""
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = font_name
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def build_docx(paragraphs: list[tuple[str, int]], output_path: Path) -> None:
    """构建 Word 文档"""
    doc = Document()
    set_docx_font(doc)

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # 添加标题页
    title = doc.add_heading("定本 柄谷行人文学论集", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("柄谷行人 著 / 陈言 译")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 添加正文内容
    for text, level in paragraphs:
        if level == 1:
            # 一级标题
            doc.add_page_break()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            # 二级标题
            doc.add_heading(text, level=2)
        elif level == 3:
            # 三级标题
            doc.add_heading(text, level=3)
        else:
            # 正文段落
            para = doc.add_paragraph()
            # 处理段落中的换行（保持段落内连贯）
            text_lines = text.split("\n")
            for i, line in enumerate(text_lines):
                run = para.add_run(line)
                run.font.size = Pt(11)
                if i < len(text_lines) - 1:
                    # 段落内换行使用软回车
                    run.add_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")
    print(f"   段落数: {len(paragraphs)}")


def main():
    print(f"📖 读取文件: {INPUT_FILE.name}")
    content = INPUT_FILE.read_text(encoding="utf-8")
    lines = content.split("\n")
    print(f"   总行数: {len(lines)}")
    print(f"   总字符数: {len(content):,}")

    # 清理内容
    cleaned = clean_content(lines)
    print(f"   清理后行数: {len(cleaned)}")

    # 合并段落
    paragraphs = merge_paragraphs(cleaned)
    title_count = sum(1 for _, level in paragraphs if level > 0)
    para_count = sum(1 for _, level in paragraphs if level == 0)
    print(f"   段落数: {para_count}")
    print(f"   标题数: {title_count}")

    # 统计标题层级
    for lvl in [1, 2, 3]:
        count = sum(1 for _, level in paragraphs if level == lvl)
        if count > 0:
            print(f"     一级标题: {count}" if lvl == 1 else
                  f"     二级标题: {count}" if lvl == 2 else
                  f"     三级标题: {count}")

    # 生成 Word 文档
    print()
    print(f"📝 生成 Word 文档...")
    build_docx(paragraphs, OUTPUT_FILE)

    # 输出文件大小
    size = OUTPUT_FILE.stat().st_size
    if size > 1024 * 1024:
        size_str = f"{size / 1024 / 1024:.1f} MB"
    else:
        size_str = f"{size / 1024:.1f} KB"
    print(f"   文件大小: {size_str}")


if __name__ == "__main__":
    main()
